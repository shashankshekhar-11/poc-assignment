from pathlib import Path
from openpyxl import Workbook
from docx import Document

DATA_DIR = Path(__file__).resolve().parents[1] / "data"

def create_products_excel():
    wb = Workbook()
    ws = wb.active
    ws.title = "Products"

    ws.append(["Product ID", "Product Name", "Category", "Price", "Stock", "Description"])

    products = [
        [101, "Laptop Pro", "Electronics", 1299.99, 15, "High performance laptop with 16GB RAM, 512GB SSD"],
        [102, "Wireless Mouse", "Accessories", 29.99, 150, "Ergonomic wireless mouse with 2.4GHz connection"],
        [103, "USB-C Cable", "Accessories", 12.99, 200, "3ft USB-C to USB-C charging and data cable"],
        [104, "Monitor 27inch", "Electronics", 349.99, 8, "4K UHD monitor with HDR support, 60Hz refresh rate"],
        [105, "Mechanical Keyboard", "Accessories", 149.99, 45, "RGB mechanical keyboard with Cherry MX switches"],
        [106, "Webcam HD", "Electronics", 79.99, 30, "1080p HD webcam with built-in microphone"],
        [107, "Headphones", "Accessories", 199.99, 60, "Noise-cancelling wireless headphones, 30hr battery"],
        [108, "Phone Stand", "Accessories", 19.99, 100, "Adjustable phone stand for desk, supports 4-7 inch"],
    ]

    for product in products:
        ws.append(product)

    wb.save(DATA_DIR / "demo-products.xlsx")
    print(f"Created: {DATA_DIR / 'demo-products.xlsx'}")

def create_policies_docx():
    doc = Document()

    doc.add_heading("Tech Store - Customer Support Guide", level=1)

    doc.add_heading("1. Shipping & Delivery", level=2)
    doc.add_paragraph(
        "We offer free shipping on all orders over $50. Standard shipping takes 5-7 business days. "
        "Express shipping (2-3 days) is available for $15. International shipping is available to most countries."
    )

    doc.add_heading("2. Return & Refund Policy", level=2)
    doc.add_paragraph(
        "You have 30 days from purchase to return any item in original condition. Electronics must be unopened. "
        "Refunds are processed within 5-7 business days. Return shipping is free for defective items."
    )

    doc.add_heading("3. Warranty Information", level=2)
    doc.add_paragraph(
        "All electronics come with a 1-year manufacturer warranty covering defects in materials and workmanship. "
        "Accessories have a 6-month warranty. Extended warranty (2-3 years) is available for purchase."
    )

    doc.add_heading("4. Customer Support Hours", level=2)
    doc.add_paragraph(
        "Monday to Friday: 9:00 AM - 6:00 PM EST\n"
        "Saturday: 10:00 AM - 4:00 PM EST\n"
        "Sunday: Closed\n"
        "Email: support@techstore.com\n"
        "Phone: 1-800-TECH-HELP"
    )

    doc.add_heading("5. Payment Methods", level=2)
    doc.add_paragraph(
        "We accept all major credit cards (Visa, Mastercard, American Express), PayPal, and Apple Pay. "
        "All payments are secure and encrypted using SSL technology."
    )

    doc.add_heading("6. Product Quality Guarantee", level=2)
    doc.add_paragraph(
        "We guarantee that all products are genuine and new. If you receive a damaged or defective item, "
        "we will replace it free of charge within 14 days of purchase."
    )

    doc.add_heading("7. Frequently Asked Questions", level=2)

    doc.add_heading("Q: Do you offer bulk discounts?", level=3)
    doc.add_paragraph("A: Yes! Orders of 10+ items receive 10% discount. Orders of 50+ receive 15% discount.")

    doc.add_heading("Q: Can I cancel my order?", level=3)
    doc.add_paragraph("A: Orders can be cancelled within 24 hours of purchase if not yet shipped.")

    doc.add_heading("Q: Do you ship internationally?", level=3)
    doc.add_paragraph("A: Yes, we ship to 150+ countries. International shipping takes 10-21 business days.")

    doc.add_heading("Q: Is my data secure?", level=3)
    doc.add_paragraph("A: Yes, we use industry-standard encryption and never share customer data with third parties.")

    doc.save(DATA_DIR / "demo-policies.docx")
    print(f"Created: {DATA_DIR / 'demo-policies.docx'}")

if __name__ == "__main__":
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    create_products_excel()
    create_policies_docx()
    print("\nDemo files created successfully!")
