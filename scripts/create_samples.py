"""Create sample xlsx and docx files for local testing."""

from pathlib import Path

from docx import Document
from openpyxl import Workbook

DATA_DIR = Path(__file__).resolve().parents[1] / "data"


def create_sample_xlsx(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sales"
    sheet.append(["Product", "Quantity", "Revenue"])
    sheet.append(["Widget A", 120, 2400])
    sheet.append(["Widget B", 85, 1700])
    sheet.append(["Gadget C", 40, 3200])
    workbook.save(path)


def create_sample_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Company Policy", level=1)
    document.add_paragraph(
        "Refund policy: Customers may request a full refund within 30 days of purchase."
    )
    document.add_paragraph(
        "Support hours: Monday to Friday, 9 AM to 6 PM IST."
    )
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Department"
    table.rows[0].cells[1].text = "Contact"
    row = table.add_row()
    row.cells[0].text = "Sales"
    row.cells[1].text = "sales@example.com"
    document.save(path)


def main() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    create_sample_xlsx(DATA_DIR / "sample.xlsx")
    create_sample_docx(DATA_DIR / "sample.docx")
    print(f"Created {DATA_DIR / 'sample.xlsx'}")
    print(f"Created {DATA_DIR / 'sample.docx'}")


if __name__ == "__main__":
    main()
