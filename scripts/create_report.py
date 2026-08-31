from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_FILE = Path(__file__).resolve().parents[1] / "RAG-POC-Report.docx"

def create_report():
    doc = Document()

    doc.add_heading("RAG POC - Implementation Report", level=0)
    doc.add_paragraph("Retrieval-Augmented Generation Proof of Concept").italic = True

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This document outlines the implementation of a Retrieval-Augmented Generation (RAG) "
        "system that allows users to upload Excel and Word documents, build a searchable vector index, "
        "and ask natural language questions about the content. The system uses FAISS for vector storage, "
        "Google Gemini for embeddings and chat, and Streamlit for the user interface."
    )

    doc.add_heading("Project Overview", level=1)
    doc.add_paragraph(
        "The RAG POC demonstrates how to build an intelligent document question-answering system. "
        "Users can upload documents, and the system will:"
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Extract text from Excel sheets and Word documents")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Split documents into semantic chunks")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Generate vector embeddings using Google Gemini")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Store embeddings in FAISS index")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Retrieve relevant chunks based on user queries")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Generate contextual answers using Gemini API")

    doc.add_heading("Architecture", level=1)
    doc.add_heading("Technology Stack", level=2)

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'

    rows_data = [
        ('UI Framework', 'Streamlit'),
        ('Vector Store', 'FAISS (IndexFlatIP)'),
        ('Embeddings', 'Google Gemini'),
        ('Chat Model', 'Google Gemini 3.6-flash'),
        ('File Parsing', 'openpyxl, python-docx'),
        ('Package Manager', 'uv (Python)')
    ]

    for i, (component, tech) in enumerate(rows_data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = component
        row_cells[1].text = tech

    doc.add_heading("System Flow", level=2)
    doc.add_paragraph(
        "1. User uploads Excel and/or Word files via Streamlit sidebar\n"
        "2. System extracts text from documents (xlsx rows, docx paragraphs, tables)\n"
        "3. Text is split into chunks (800 chars with 150 char overlap)\n"
        "4. Each chunk is embedded using Gemini API\n"
        "5. Embeddings stored in FAISS IndexFlatIP\n"
        "6. User asks a question\n"
        "7. Question is embedded and searched in FAISS\n"
        "8. Top-K (4) results retrieved with relevance scores\n"
        "9. Results sent to Gemini with question for context-aware answering\n"
        "10. Answer displayed with source references"
    )

    doc.add_heading("Implementation Details", level=1)

    doc.add_heading("File Structure", level=2)
    doc.add_paragraph(
        "src/rag_poc/\n"
        "├── app.py              Streamlit UI application\n"
        "├── config.py           Configuration and API setup\n"
        "├── loaders.py          Excel/Word file parsing\n"
        "├── chunking.py         Text segmentation logic\n"
        "├── embeddings.py       Gemini embedding wrapper\n"
        "├── vectorstore.py      FAISS index management\n"
        "└── rag.py              Q&A and retrieval logic"
    )

    doc.add_heading("Key Features", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Multi-document support").bold = True
    p.add_run(" - Process both Excel and Word documents simultaneously")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Smart chunking").bold = True
    p.add_run(" - Merges related content and respects document boundaries")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Vector normalization").bold = True
    p.add_run(" - Uses cosine similarity for accurate relevance scoring")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Persistent storage").bold = True
    p.add_run(" - FAISS index and metadata saved to disk")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Conversational UI").bold = True
    p.add_run(" - Handles casual greetings and document-based questions")

    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Source attribution").bold = True
    p.add_run(" - Shows which documents and sections informed the answer")

    doc.add_heading("Demo Walkthrough", level=1)

    doc.add_heading("Step 1: File Upload", level=2)
    doc.add_paragraph("Screenshot: [INSERT - Streamlit sidebar with file uploaders]")
    doc.add_paragraph(
        "Users select an Excel file (.xlsx) and/or Word file (.docx) from their computer. "
        "The app shows file size and name."
    )

    doc.add_heading("Step 2: Build Index", level=2)
    doc.add_paragraph("Screenshot: [INSERT - Build Index button and progress]")
    doc.add_paragraph(
        "Clicking 'Build Index' initiates the pipeline:\n"
        "• Files are parsed for text content\n"
        "• Text is chunked into semantic pieces\n"
        "• Embeddings are generated via Gemini API\n"
        "• FAISS index is built and saved\n"
        "Status shows chunk count (e.g., 'Index ready with 30 chunks')"
    )

    doc.add_heading("Step 3: Ask Questions", level=2)
    doc.add_paragraph("Screenshot: [INSERT - Chat interface with question and answer]")
    doc.add_paragraph(
        "Users type natural language questions in the chat box. Examples:\n"
        "• 'What products do you have?'\n"
        "• 'What is the refund policy?'\n"
        "• 'How much does a laptop cost?'\n"
        "• 'What are your support hours?'"
    )

    doc.add_heading("Step 4: View Results", level=2)
    doc.add_paragraph("Screenshot: [INSERT - Answer with sources expanded]")
    doc.add_paragraph(
        "The system displays:\n"
        "• The AI-generated answer based on document content\n"
        "• Source chunks (expandable) showing which documents were used\n"
        "• Relevance scores (0.0-1.0) for each source"
    )

    doc.add_heading("Sample Output", level=1)

    doc.add_heading("Query: 'What is the return policy?'", level=2)
    doc.add_paragraph("Sources Retrieved: 2", style='List Number')
    doc.add_paragraph(
        "[1] demo-policies.docx (Paragraph 2)\n"
        "Score: 0.892\n"
        "'You have 30 days from purchase to return any item in original condition...'"
    )
    doc.add_paragraph(
        "[2] demo-policies.docx (Paragraph 3)\n"
        "Score: 0.756\n"
        "'All electronics come with a 1-year manufacturer warranty...'"
    )

    doc.add_heading("Answer", level=3)
    doc.add_paragraph(
        "Based on the provided context, you have 30 days from purchase to return any item "
        "in its original condition. Electronics must be unopened. Refunds are processed within "
        "5-7 business days, and return shipping is free for defective items."
    )

    doc.add_heading("Performance Metrics", level=1)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'

    metrics = [
        ('Chunk Size', '800 characters'),
        ('Embedding Model', 'Gemini Embedding 001'),
        ('Index Type', 'FAISS IndexFlatIP (cosine similarity)'),
        ('Retrieved Results', 'Top-4 most relevant chunks')
    ]

    for i, (metric, value) in enumerate(metrics, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = metric
        row_cells[1].text = value

    doc.add_heading("How to Use", level=1)
    doc.add_heading("Installation", level=2)
    doc.add_paragraph(
        "1. Clone the repository\n"
        "2. Run: uv sync\n"
        "3. Copy .env.example to .env\n"
        "4. Add your GEMINI_API_KEY to .env"
    )

    doc.add_heading("Running the App", level=2)
    doc.add_paragraph("uv run streamlit run src/rag_poc/app.py")
    doc.add_paragraph("Access at: http://localhost:8501")

    doc.add_heading("Limitations & Future Work", level=1)

    doc.add_heading("Current Limitations", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Single session - no multi-user support")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Fixed chunk size - 800 characters (not adaptive)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("No authentication - suitable for internal use")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Limited file formats - Excel and Word only")

    doc.add_heading("Potential Enhancements", level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("PDF support")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Hybrid search (lexical + semantic)")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Query expansion and re-ranking")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Multi-turn conversation history")
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Database persistence across sessions")

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "This RAG POC demonstrates a practical implementation of retrieval-augmented generation "
        "using modern APIs and open-source libraries. The system successfully combines document "
        "processing, vector embeddings, and language models to provide accurate, context-grounded "
        "answers to user queries."
    )

    doc.add_paragraph(
        "The simple, straightforward codebase makes it easy to understand, extend, and deploy "
        "for real-world document Q&A applications."
    )

    doc.save(OUTPUT_FILE)
    print(f"Report created: {OUTPUT_FILE}")

if __name__ == "__main__":
    create_report()
