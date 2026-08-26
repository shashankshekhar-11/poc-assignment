from openpyxl import load_workbook
from docx import Document

def load_excel_file(file_path):
    all_rows = []

    workbook = load_workbook(file_path, read_only=True, data_only=True)

    for sheet in workbook.worksheets:
        sheet_name = sheet.title
        rows = sheet.iter_rows(values_only=True)

        header = next(rows, None)

        for row_number, row_values in enumerate(rows, start=2):
            row_text = f"Sheet: {sheet_name} | "

            for col_number, cell_value in enumerate(row_values):
                if cell_value is not None and str(cell_value).strip() != "":
                    if header and col_number < len(header):
                        column_name = header[col_number]
                    else:
                        column_name = f"Column{col_number + 1}"
                    row_text += f"{column_name}: {cell_value} | "

            row_data = {
                "text": row_text,
                "file": "sample.xlsx",
                "location": f"Sheet {sheet_name}, Row {row_number}"
            }
            all_rows.append(row_data)

    workbook.close()
    return all_rows

def load_word_file(file_path):
    all_content = []

    doc = Document(file_path)

    for para_number, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            para_data = {
                "text": text,
                "file": "sample.docx",
                "location": f"Paragraph {para_number}"
            }
            all_content.append(para_data)

    for table_number, table in enumerate(doc.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            cells_text = []
            for cell in row.cells:
                if cell.text.strip():
                    cells_text.append(cell.text.strip())

            if cells_text:
                table_data = {
                    "text": " | ".join(cells_text),
                    "file": "sample.docx",
                    "location": f"Table {table_number}, Row {row_number}"
                }
                all_content.append(table_data)

    return all_content

def load_files(excel_path=None, word_path=None):
    all_data = []

    if excel_path:
        excel_data = load_excel_file(excel_path)
        all_data.extend(excel_data)

    if word_path:
        word_data = load_word_file(word_path)
        all_data.extend(word_data)

    if not all_data:
        raise ValueError("No data found in files")

    return all_data
